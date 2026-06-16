from flask import render_template, redirect, url_for, flash, request, session, abort, Blueprint, send_file, send_from_directory, current_app, make_response, Response, jsonify
from flask_login import login_user, login_required, logout_user, current_user
from flask_wtf import FlaskForm
from werkzeug.utils import secure_filename
import pandas as pd
from openpyxl import Workbook
from io import BytesIO
from forms import SignupForm, LoginForm, EditProfileForm, ProductForm, AddToCartForm, SettingsForm, CommissionPlanForm
from models import Admin, User, Buyer, Seller, Product, Cart, ProductComponent, ProductImage, Order, Subscription, CommissionPlan
from database import db
from functools import wraps
from fpdf import FPDF
from docx import Document
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter
import io
import qrcode
from qrcode.constants import ERROR_CORRECT_H
import os

admin_bp = Blueprint('admin_routes', __name__, static_folder='static', static_url_path='/static')

#----------------FUNCTIONS DECORATOR-------------------
# Custom decorator for role-based access control
def role_required(*roles):
    def wrapper(func):
        @wraps(func)
        def decorated_view(*args, **kwargs):
            if current_user.role not in roles:
                flash('Access denied. You do not have permission to access this page.', 'danger')
                return redirect(url_for('routes.index'))
            return func(*args, **kwargs)
        return decorated_view
    return wrapper

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in current_app.config['ALLOWED_EXTENSIONS']

#---------ADMIN-----------
@admin_bp.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    form = LoginForm()
    
    # Get the 'next' parameter, default to the admin dashboard
    next_page = request.args.get('next', url_for('admin_routes.admin_dashboard'))

    # Ensure next_page is a relative URL to prevent open redirects
    if next_page.startswith('/') is False:
        next_page = url_for('admin.admin_dashboard')

    if form.validate_on_submit():
        admin = Admin.query.filter_by(email=form.email.data).first()
        
        if admin and admin.check_password(form.password.data):
            login_user(admin)
            flash('Admin login successful.', 'success')

            print(f"Redirecting to: {next_page}")  # Debugging
            
            return redirect(next_page)

        flash('Invalid email or password.', 'danger')

    return render_template('login.html', title='Admin Login', form=form, next=next_page)

# Update the admin_dashboard route
@admin_bp.route('/admin_dashboard')
@login_required
@role_required('admin')
def admin_dashboard():
    total_users = User.query.count()
    total_buyers = Buyer.query.count()
    total_sellers = Seller.query.count()
    total_admins = User.query.filter_by(role='admin').count()
    products = Product.query.filter_by(status='Pending').all()
    
    print(f"Fetched {len(products)} pending products")  # Debugging statement
    
    admin = Admin.query.filter(Admin.role != 'admin').all()
    return render_template('admin_dashboard.html', 
                           total_users=total_users, 
                           total_buyers=total_buyers, 
                           total_sellers=total_sellers, 
                           total_admins=total_admins, 
                           products=products,
                           admin=admin)

@admin_bp.route('/admin/view_user/<int:user_id>', methods=['GET'])
@login_required
def admin_view_user(user_id):
    user = User.query.get_or_404(user_id)
    # Ensure profile image URL uses forward slashes
    profile_image_url = user.profile_image.replace('\\', '/')    
    current_app.logger.debug(f"Profile image path in admin view: {profile_image_url}")
    return render_template('view_user.html', user=user, profile_image_url=profile_image_url)

@admin_bp.route('/admin/edit_user/<int:user_id>', methods=['GET', 'POST'])
@login_required
def admin_edit_user(user_id):
    user = User.query.get_or_404(user_id)
    form = EditProfileForm()  # Instantiate the form

    if request.method == 'POST':
        if form.validate_on_submit():
            # Logic for editing user details, e.g., saving the image if uploaded
            user.username = form.username.data
            user.email = form.email.data
            
            # Check for the uploaded profile image
            if form.profile_image.data:
                # Save the image logic here
                file = form.profile_image.data
                filename = secure_filename(file.filename)
                file.save(os.path.join(admin_bp.root_path, 'static/users', f"{user.username}_{user.id}/profile_images", filename))
                user.profile_image = f"users/{user.username}_{user.id}/profile_images/{filename}"

            db.session.commit()  # Commit changes to the database
            return redirect(url_for('admin_routes.admin_dashboard'))

    # Populate the form with existing user data
    form.username.data = user.username
    form.email.data = user.email

    return render_template('edit_user.html', form=form, user=user)

@admin_bp.route('/admin/delete_user/<int:user_id>', methods=['GET', 'POST'])
def admin_delete_user(user_id):
    user = User.query.get_or_404(user_id)
    if request.method == 'POST':
        # Redirect to confirm delete page
        return redirect(url_for('routes.admin_confirm_delete_user', user_id=user.id))
    return render_template('delete_user.html', user=user)

@admin_bp.route('/admin/confirm_delete_user/<int:user_id>', methods=['POST'])
def admin_confirm_delete_user(user_id):
    user = User.query.get_or_404(user_id)
    try:
        db.session.delete(user)
        db.session.commit()
        flash('User has been deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash('Error deleting user: ' + str(e), 'danger')
    return redirect(url_for('routes.admin_dashboard'))

@admin_bp.route('/preview_users_pdf')
def preview_users_pdf():
    # Create a BytesIO buffer
    buffer = BytesIO()    
    # Create a PDF canvas
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    # Set title
    p.setFont("Helvetica-Bold", 16)
    p.drawString(200, height - 50, "User List")
    # Set column titles
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, height - 100, "ID")
    p.drawString(150, height - 100, "Username")
    p.drawString(300, height - 100, "Email")
    p.drawString(450, height - 100, "Role")
    # Add a line for the header
    p.line(50, height - 105, 550, height - 105)
    # Reset font for data
    p.setFont("Helvetica", 12)
    # Sample data retrieval, replace with your actual database call
    users = User.query.all()  # Fetch all users from the database
    # Print user data in the table
    y = height - 120  # Starting y position for user data
    for user in users:
        p.drawString(50, y, str(user.id))
        p.drawString(150, y, user.username)
        p.drawString(300, y, user.email)
        p.drawString(450, y, user.role)
        y -= 20  # Move down for the next row
    # Finalize the PDF
    p.showPage()
    p.save()    
    # Move the buffer cursor to the beginning
    buffer.seek(0)
    # Return the PDF as a response to display in the browser
    return Response(buffer, mimetype='application/pdf')

@admin_bp.route('/export_users_pdf')
def export_users_pdf():
    # Your PDF generation code
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    # Add content to PDF
    p.showPage()
    p.save()
    buffer.seek(0)
    return Response(buffer, mimetype='application/pdf',
                    headers={"Content-Disposition": "attachment;filename=users.pdf"})

@admin_bp.route('/export_users_excel')
@login_required
@role_required('admin')
def export_users_excel():
    # Assuming you have a list of users fetched from the database
    users = User.query.filter(User.role != 'admin').all()
    # Create a DataFrame
    data = {
        'ID': [user.id for user in users],
        'Username': [user.username for user in users],
        'Email': [user.email for user in users],
        'Role': [user.role for user in users]
    }
    df = pd.DataFrame(data)
    # Create an output buffer
    output = io.BytesIO()
    # Using 'openpyxl' to create an Excel file
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Users')
    output.seek(0)  # Go to the beginning of the BytesIO buffer
    # Use 'download_name' instead of 'attachment_filename'
    return send_file(output, download_name='users.xlsx', as_attachment=True)

@admin_bp.route('/export_users_word')
@login_required
@role_required('admin')
def export_users_word():
    # Fetch users from the database
    users = User.query.filter(User.role != 'admin').all()
    # Create a Word Document
    doc = Document()
    doc.add_heading('User List', 0)
    # Add a table
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Username'
    hdr_cells[2].text = 'Email'
    hdr_cells[3].text = 'Role'
    # Populate the table with user data
    for user in users:
        row_cells = table.add_row().cells
        row_cells[0].text = str(user.id)
        row_cells[1].text = user.username
        row_cells[2].text = user.email
        row_cells[3].text = user.role
    # Save to a BytesIO object
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)  # Move to the beginning of the BytesIO object
    # Use 'download_name' instead of 'attachment_filename'
    return send_file(output, download_name='users.docx', as_attachment=True)

@admin_bp.route('/admin/products')
@login_required
@role_required('admin')
def admin_products():
    products = Product.query.all()
    return render_template('admin_products.html', products=products)

@admin_bp.route('/admin/view_product/<int:product_id>', methods=['GET'])
def admin_view_product(product_id):
    # Query the product by ID
    product = Product.query.get_or_404(product_id)    
    # Render the product details template with the product information
    return render_template('admin_view_product.html', product=product)

# Define the folder for saving QR codes
QRCODE_FOLDER = 'static/qr_codes'
if not os.path.exists(QRCODE_FOLDER):
    os.makedirs(QRCODE_FOLDER)


@admin_bp.route('/admin/approve_product/<int:product_id>', methods=['POST'])
@login_required
def approve_product(product_id):
    if current_user.role != 'admin':
        flash('Unauthorized access', 'danger')
        return redirect(url_for('admin_routes.admin_dashboard'))
    
    product = Product.query.get_or_404(product_id)
    product.status = 'Approved'

    # Generate QR Code with version 2
    qr_data = f"Product ID: {product.id}, Product Name: {product.name}"
    qr = qrcode.QRCode(
        version=2,  # Specifies QR Code version
        error_correction=ERROR_CORRECT_H,  # High error correction level
        box_size=10,
        border=4,
    )
    qr.add_data(qr_data)
    qr.make(fit=True)

    # Save the QR code to a file
    filename = f"product_{product.id}_qr.png"
    qr_image_path = os.path.join(QRCODE_FOLDER, filename)
    try:
        # Save QR Code to a file
        img = qr.make_image(fill_color="black", back_color="white")
        with open(qr_image_path, 'wb') as img_file:
            img.save(img_file)
        
        # Store public QR path in the database
        public_qr_path = url_for('static', filename=f'qr_codes/{filename}')
        product.qr_code = public_qr_path
        db.session.commit()

        flash(f'Product "{product.name}" has been approved and QR code generated.', 'success')
    except Exception as e:
        flash(f'Failed to generate QR code: {str(e)}', 'danger')

    return redirect(url_for('admin_routes.admin_dashboard'))

@admin_bp.route('/admin/reject_product/<int:product_id>', methods=['POST'])
@login_required
def reject_product(product_id):
    product = Product.query.get_or_404(product_id)
    product.status = 'Rejected'
    db.session.commit()
    flash('Product rejected.', 'success')
    return redirect(url_for('routes.admin_dashboard'))

@admin_bp.route('/admin/products/edit/<int:product_id>', methods=['GET', 'POST'])
@login_required
def edit_product(product_id):
    product = Product.query.get_or_404(product_id)
    user = current_user
    # Check if the user has the necessary permissions
    if user.role == 'admin' or (user.role == 'seller' and user.id == product.seller_id):
        form = ProductForm()
        if request.method == 'POST':
            form = ProductForm(request.form, obj=product)
            if form.validate_on_submit():
                form.populate_obj(product)
                if form.image.data:
                    image_file = form.image.data
                    if allowed_file(image_file.filename):
                        user_folder = os.path.join('static', 'users', f"{current_user.username}_{current_user.id}", 'product_images')
                        os.makedirs(user_folder, exist_ok=True)
                        filename = secure_filename(image_file.filename)
                        file_path = os.path.join(user_folder, filename)
                        image_file.save(file_path)
                        product.image_url = os.path.join(f'users/{current_user.username}_{current_user.id}/product_images', filename)
                db.session.commit()
                flash('Product updated successfully!', 'success')
                return redirect(url_for('routes.admin_products') if user.role == 'admin' else url_for('seller_dashboard'))
        elif request.method == 'GET':
            form.name.data = product.name
            form.description.data = product.description
            form.price.data = product.price
            form.category.data = product.category
        return render_template('edit_product.html', form=form, product=product)
    else:
        flash('You do not have permission to edit this product.', 'danger')
        return redirect(url_for('routes.admin_products'))

@admin_bp.route('/admin/products/delete/<int:product_id>', methods=['POST'])
@login_required
def admin_delete_product(product_id):
    product = Product.query.get_or_404(product_id)
    user = current_user    
    if user.role == 'admin' or (user.role == 'seller' and user.id == product.seller_id):
        db.session.delete(product)
        db.session.commit()
        flash('Product deleted successfully!', 'success')
    else:
        flash('You do not have permission to delete this product.', 'danger')
    return redirect(url_for('routes.admin_products') if user.role == 'admin' else url_for('seller_dashboard'))

#-------------------COMMISSION--------------

#--------------ORDERS -------------------
@admin_bp.route('/orders', methods=['GET'])
@login_required
def admin_orders():
    if current_user.role != 'admin':
        flash("You are not authorized to view this page.", "danger")
        return redirect(url_for('routes.marketplace'))

    # Fetch all approved orders
    orders = Order.query.filter_by(status="Approved").all()

    return render_template('admin_dashboard.html', orders=orders)

@admin_bp.route('/static/<path:filename>')
def static_files(filename):
    return send_from_directory(admin_bp.static_folder, filename)

# Admin view for managing subscriptions
@admin_bp.route('/admin/subscriptions', methods=['GET', 'POST'])
@login_required
def manage_subscriptions():
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for('routes.index'))

    subscriptions = Subscription.query.all()
    if request.method == 'POST':
        # Add a new subscription offer
        name = request.form.get('name')
        description = request.form.get('description')
        price = float(request.form.get('price'))
        subscription = Subscription(name=name, description=description, price=price)
        db.session.add(subscription)
        db.session.commit()
        flash("Subscription added successfully.", "success")
        return redirect(url_for('admin_routes.manage_subscriptions'))

    return render_template('admin_subscriptions.html', subscriptions=subscriptions)

# Toggle subscription status
@admin_bp.route('/admin/subscriptions/<int:subscription_id>/toggle', methods=['POST'])
@login_required
def toggle_subscription(subscription_id):
    if not current_user.is_admin:
        flash("Access denied.", "danger")
        return redirect(url_for('index'))

    subscription = Subscription.query.get_or_404(subscription_id)
    subscription.status = 'active' if subscription.status == 'inactive' else 'inactive'
    db.session.commit()
    flash("Subscription status updated.", "success")
    return redirect(url_for('admin_routes.manage_subscriptions'))